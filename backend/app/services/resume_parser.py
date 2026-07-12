"""
Extracts plain text from uploaded resume files (PDF or DOCX).
Keeping this isolated makes it easy to swap in OCR (for scanned PDFs)
later without touching the rest of the pipeline.
"""
import io
from pathlib import Path

import pdfplumber
from docx import Document


class UnsupportedFileType(Exception):
    pass


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Given raw file bytes and the original filename (used to detect type),
    returns extracted plain text. Raises UnsupportedFileType for anything
    other than .pdf and .docx.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _extract_from_pdf(file_bytes)
    elif suffix == ".docx":
        return _extract_from_docx(file_bytes)
    else:
        raise UnsupportedFileType(
            f"Unsupported file type '{suffix}'. Please upload a PDF or DOCX resume."
        )


def _extract_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    text = "\n".join(text_parts).strip()

    if not text:
        # Likely a scanned/image-only PDF. Flag it clearly rather than
        # silently returning empty text (OCR fallback is a Phase-2+ add-on).
        raise ValueError(
            "No extractable text found in PDF. It may be a scanned image — "
            "OCR support isn't wired up yet in this phase."
        )
    return text


def _extract_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text out of tables (many resumes use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    text = "\n".join(paragraphs).strip()
    if not text:
        raise ValueError("No extractable text found in DOCX file.")
    return text
